'''江都人民医院'''
import os
import re
import json
import glob

from docx import Document


def get_basic_info(text):
    '''从封面的文本内容获取个人基本信息'''
    name, sex, age, num, date = re.findall(r'姓    名(.*)性    别(.*)年    龄(.*)体检编号(.*)体检次数.*预约日期(\d+-\d+-\d+)', text)[0]
    return {
        '姓名': name,
        '性别': sex,
        '年龄': age,
        '体检编号': num,
        '预约日期': date
    }



def walk_inner_tables(cell, samples):
    '''体检指标数据位于table中的cell嵌套的table中'''
    for table_ in cell.tables:
        for cell_ in table_._cells:
            cell_header = table_._cells[0].text.strip()
            if cell_.text.strip() == '小结':# 舍弃彩超等部分中的小结部分
                continue
            for table__ in cell_.tables:
                s = []
                for cell__ in table__._cells:
                    text__ = cell__.text.strip()
                    s.append(text__)
                if s and {cell_header: s} not in samples and '【检查所见】' not in s:
                    samples.append({cell_header: s})
    return samples


def generate_standard_data(raw_samples_darta):
    '''将从.docx中获取的原始指标数据转化为标准格式的数据'''
    std_samples = []
    for group_samples in raw_samples_darta:
        # group_samples: {"基本信息": ["身高","157","收缩压","164"]}
        for k, v in group_samples.items():
            if len(v) > 1 and len(v) % 2 == 0:
                ref_range, unit = None, None
                if v[0] == '项目' and v[1] == '检测值':
                    sample_dict = dict(zip(v[10::5], v[11::5]))
                    ref_range = v[13::5]
                    unit = v[14::5]
                else:
                    sample_dict = dict(zip(v[::2], v[1::2]))
                for i, (item_name, item_value) in enumerate(sample_dict.items()):
                    if item_name:
                        std_s = {
                            'group_name': k,
                            'item_name': item_name,
                            'value': item_value
                        }
                        if ref_range and ref_range[i]:
                            std_s['refrange'] = ref_range[i]
                        if unit and unit[i]:
                            std_s['unit'] = unit[i]
                        std_samples.append(std_s)
    return std_samples


def extract_data(document, file_name):
    res = {}
    summary = ''
    raw_samples = []
    res['file_name'] = file_name
    for i, table in enumerate(document.tables):
        temp_paragraph = ''
        samples = []
        for cell in table._cells:
            text = cell.text
            if i in [1, 2]:# 综述部分
                summary_ = text.split('  ')[0]
                try:
                    int(summary_)
                    summary += f"{summary_} "
                except ValueError:
                    if summary_:
                        summary += f"{summary_}\n"
            temp_paragraph += text
            samples = walk_inner_tables(cell, samples)
        if samples:
            raw_samples.extend(samples)
        if i == 0:# 封面--> 提取个人信息
            res['basic_info'] = get_basic_info(temp_paragraph)

    res['summary'] = summary.replace(
        '\n疾病诊断\n疾病诊断', ''
    ).replace(
        '\n详情请到以上相关专科进行咨询、治疗或随访。', ''
    ).replace(
        '\n其他阳性发现\n其他阳性发现', ''
    )
    std_samples = generate_standard_data(raw_samples)
    res['samples'] = std_samples
    return res


def main(path):
    res = []
    docx_paths = glob.glob(path)
    for docx_path in docx_paths:
        document = Document(docx_path)
        data = extract_data(document, os.path.basename(docx_path))
        res.append(data)

    with open('/Users/har/Desktop/江都std.json', 'w') as f:
        json.dump(res, f, ensure_ascii=False, indent=2)

path = '/Users/har/Desktop/江都人民医院（128人-个人需要删信息WORD)docx/*.docx'
main(path)
